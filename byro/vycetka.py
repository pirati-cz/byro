#!/usr/bin/env python3
# -*- encoding: utf-8 -*-

import datetime
from .bredmine import Utils
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from byro.bredmine import BRedmine


class Output:
	def __init__(self, data):
		self.data = data

	def show(self):
		pass


class Cmd(Output):
	""" Exports RedmineData into command line """

	def show(self):
		for ref in self.data.refunds:
			print("\n", self.data.refunds[ref]['fullname'])
			for issue in self.data.refunds[ref]['data']:
				self._print_time_entry(issue)

		print("\nCelkem:", self.data.totalSum, "h",
		      "v měsící", self.data.month, self.data.year,
		      ", uživatel", self.data.user,
		      "v projektu", self.data.project
		      )

	@staticmethod
	def _print_time_entry(te):
		print(te.issue, "\t", te.spent_on, te.hours, te.activity, "\t", te.comments)


class DocX(Output):
	""" Exports RedmineData into docx files """

	def __init__(self, data, filename=None):
		Output.__init__(self, data)
		if filename is None:
			self.filename = "vycetka-" + data.month + "-" + data.user.lastname + ".docx"
		else:
			self.filename = filename

	def show(self):
		self._createDoc()

	def _createDoc(self, date=None):
		"""
		:param date:
		:type date: None
		:type date: datetime
		:return:
		"""
		if date is None:
			date = datetime.datetime.now()

		document = Document()
		styles = document.styles
		self.table_style = styles['TableGrid']

		h1 = document.add_heading('Výčetka', level=1)

		h1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

		p1 = document.add_paragraph(
			'pro výpočet náhrady mzdy nebo výdělku ušlého v souvislosti s výkonem funkce neuvolněného člena Zastupitelstva hlavního města Prahy')
		p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

		p2 = document.add_paragraph('za měsíc ' + self.data.month + ' ' + str(self.data.year))
		p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

		p2 = document.add_paragraph(text='Jméno a příjmení: ' + str(self.data.user))

		self._createTable(document)

		document.add_paragraph('Prohlašuji, že výše uvedené údaje jsou pravdivé.')

		document.add_paragraph(
			'Datum:	' + date.strftime('%d. %m. %Y') + '					Podpis:  	' + str(self.data.user))

		document.save(self.filename)

		print("Výčetka vygenerována do souboru %s" % self.filename)

	def _createTable(self, document):
		"""
		:param document: document main class
		:type document: docx.document.Document
		:return:
		"""
		self._createTableHeading(document)
		for refund in self.data.refunds:
			firstRow = None
			table = document.add_table(0, 6)
			table.style = self.table_style
			for event in self.data.refunds[refund]['data']:
				r = table.add_row()
				if firstRow is None:
					r.cells[0].text = self.data.refunds[refund]['fullname']
					firstRow = r
				else:
					firstRow.cells[0].merge(r.cells[0])
				comment = Utils.pick_time(event.comments)
				r.cells[1].text = event.spent_on.strftime("%d. %m. %Y")
				r.cells[2].text = comment[0]
				r.cells[3].text = str(event.hours)
				r.cells[4].text = str(event.hours)
				r.cells[5].text = comment[1]

		t = document.add_table(0, 6)
		r = t.add_row()
		r.cells[0].text = "Celkem"
		r.cells[3].text = str(self.data.totalSum)
		r.cells[4].text = str(self.data.totalSum)

	def _createTableHeading(self, document):
		table = document.add_table(2, 6)
		table.style = self.table_style

		r1 = table.row_cells(0)
		r2 = table.row_cells(1)

		r1[0].merge(r2[0])
		r1[0].text = 'Výkon funkce'

		r1[1].merge(r2[1])
		r1[1].text = 'Datum'
		r1[2].merge(r2[2])
		r1[2].text = 'Hodiny od - do'
		r1[3].merge(r1[4])
		r1[3].text = 'Počet hodin'
		r1[5].merge(r2[5])
		r1[5].text = 'Poznámka'
		r2[3].text = 'celkem'
		r2[4].text = 'k náhradě'


def vycetka_wrapper(args):
	redmine = BRedmine(args.user, args.url, args.project, month=args.month)
	data = redmine.get_data()
	# data.export_to_bin_file("data-6-OP.p")
	# cmd = Cmd(data)
	docx = DocX(data, args.out)
	docx.show()

	# cmd.show()


if __name__ == "__main__":
	pass